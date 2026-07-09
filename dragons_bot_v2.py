#!/usr/bin/env python3
import os, re, logging, json
from datetime import datetime
from telegram import Update
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, filters, ContextTypes, ConversationHandler
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

BOT_TOKEN = os.environ.get("BOT_TOKEN")
ADMIN_ID = int(os.environ.get("ADMIN_ID", "0"))

logging.basicConfig(format='%(asctime)s - %(levelname)s - %(message)s', level=logging.INFO)

# ====== حفظ الداتا في Railway Variables ======
def load_data():
    codes_str = os.environ.get("VALID_CODES", "{}")
    users_str = os.environ.get("AUTHORIZED_USERS", "{}")
    try:
        codes = json.loads(codes_str)
        users = {int(k): v for k, v in json.loads(users_str).items()}
    except:
        codes, users = {}, {}
    return codes, users

def save_data(codes, users):
    # بنحفظ في ملف مؤقت + بنطبع للـ logs عشان تقدر تشوفه
    try:
        with open("data.json", "w") as f:
            json.dump({"codes": codes, "users": {str(k): v for k, v in users.items()}}, f)
    except:
        pass

def load_from_file():
    if os.path.exists("data.json"):
        try:
            with open("data.json") as f:
                data = json.load(f)
                codes = data.get("codes", {})
                users = {int(k): v for k, v in data.get("users", {}).items()}
                return codes, users
        except:
            pass
    return load_data()

valid_codes, authorized_users = load_from_file()
user_templates = {}

WAITING_CODE = 1
WAITING_TEMPLATE_CHOICE = 2

# ====== حقول الفواتير ======
DRAGONS_FIELDS = ['الاسم', 'الرقم', 'العنوان', 'اللون', 'العرض', 'المقاس', 'العدد', 'الاجمالي']
NF_FIELDS = ['اسم العميل', 'اكونت العميل', 'رقم الفاتورة', 'اسم الصفحة', 'اسم IT', 'العنوان', 'الرقم', 'الألوان', 'العدد', 'المقاس', 'نوع المنتج', 'سعر المنتج', 'مصاريف الشحن', 'الاجمالي']
MM_FIELDS = ['اسم الاكونت', 'IT', 'الاسم', 'العنوان', 'موبايل 1', 'موبايل 2', 'المنتج', 'الالوان', 'المقاس', 'الكمية', 'السعر', 'الشحن', 'الاجمالي']
TEMPLATES = {'dragon': DRAGONS_FIELDS, 'fashion': NF_FIELDS, 'mm': MM_FIELDS}

# ====== مساعد Word ======
def set_border(cell, color='000000', sz='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:color'), color)
        tcPr.append(b)

def set_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def cell_text(cell, text, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT, color=None):
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(str(text or ''))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

# ====== Dragon.s ======
def make_dragons_invoice(doc, inv):
    NOTE = 'في حالة عدم استلام الاوردر يتم دفع مصريف الشحن كامله للمندوب.'
    table = doc.add_table(rows=0, cols=4)
    table.style = 'Table Grid'
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[3])
    set_border(cell); set_bg(cell, 'D9D9D9')
    cell_text(cell, 'Dragon.s', bold=True, size=12)
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[3])
    set_border(cell); cell_text(cell, 'التاريخ/ رقم الاوردر/', bold=True)
    for label, field in [('الاسم:', 'الاسم'), ('الرقم', 'الرقم'), ('العنوان', 'العنوان')]:
        row = table.add_row()
        set_border(row.cells[0]); set_bg(row.cells[0], 'D9D9D9'); cell_text(row.cells[0], label)
        cell = row.cells[1].merge(row.cells[3]); set_border(cell); cell_text(cell, inv.get(field, ''))
    row = table.add_row()
    set_border(row.cells[0]); set_bg(row.cells[0], 'D9D9D9'); cell_text(row.cells[0], 'اللون:')
    set_border(row.cells[1]); cell_text(row.cells[1], inv.get('اللون', ''))
    set_border(row.cells[2]); set_bg(row.cells[2], 'D9D9D9'); cell_text(row.cells[2], 'العرض:')
    set_border(row.cells[3]); cell_text(row.cells[3], inv.get('العرض', ''))
    row = table.add_row()
    set_border(row.cells[0]); set_bg(row.cells[0], 'D9D9D9'); cell_text(row.cells[0], 'المقاس')
    set_border(row.cells[1]); cell_text(row.cells[1], inv.get('المقاس', ''))
    set_border(row.cells[2]); set_bg(row.cells[2], 'D9D9D9'); cell_text(row.cells[2], 'العدد:')
    set_border(row.cells[3]); cell_text(row.cells[3], inv.get('العدد', ''))
    row = table.add_row()
    set_border(row.cells[0]); set_bg(row.cells[0], 'D9D9D9'); cell_text(row.cells[0], 'الاجمالي')
    set_border(row.cells[1]); cell_text(row.cells[1], inv.get('الاجمالي', ''))
    set_border(row.cells[2]); set_bg(row.cells[2], 'D9D9D9'); cell_text(row.cells[2], 'IT :')
    set_border(row.cells[3]); cell_text(row.cells[3], '')
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[3]); set_border(cell)
    cell_text(cell, NOTE, bold=True, size=9)

# ====== New Fashion ======
def make_nf_invoice(doc, inv):
    NOTE = 'عزيزي العميل يحق لك قبل الاستلام معاينة المنتج والتأكد من المقاس ومن جودة المنتج  في حاله عدم الاستلام برجاء دفع مصاريف الشحن للمندوب    الاستبدال حد اقصي ثلاث ايام من تاريخ التسليم مع دفع مصاريف الشحن مره اخري ولا يوجد استرجاع'
    today = datetime.now().strftime('%d / %m / %Y')
    table = doc.add_table(rows=0, cols=6)
    table.style = 'Table Grid'
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[5])
    set_border(cell); set_bg(cell, '1F3864')
    cell_text(cell, '--NEW FASHION COMPANY--', bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color='FFFFFF')
    row = table.add_row()
    set_border(row.cells[0]); cell_text(row.cells[0], inv.get('رقم الفاتورة', ''))
    set_border(row.cells[1]); set_bg(row.cells[1], 'D9D9D9'); cell_text(row.cells[1], 'رقم الفاتورة')
    set_border(row.cells[2]); cell_text(row.cells[2], inv.get('اسم IT', ''))
    set_border(row.cells[3]); set_bg(row.cells[3], 'D9D9D9'); cell_text(row.cells[3], 'اسم الـ IT')
    set_border(row.cells[4]); cell_text(row.cells[4], inv.get('اسم العميل', ''))
    set_border(row.cells[5]); set_bg(row.cells[5], 'D9D9D9'); cell_text(row.cells[5], 'اسم العميل')
    row = table.add_row()
    set_border(row.cells[0]); cell_text(row.cells[0], today)
    set_border(row.cells[1]); set_bg(row.cells[1], 'D9D9D9'); cell_text(row.cells[1], 'تاريخ الفاتورة')
    set_border(row.cells[2]); cell_text(row.cells[2], inv.get('اسم الصفحة', ''))
    set_border(row.cells[3]); set_bg(row.cells[3], 'D9D9D9'); cell_text(row.cells[3], 'اسم الصفحة')
    set_border(row.cells[4]); cell_text(row.cells[4], inv.get('اكونت العميل', ''))
    set_border(row.cells[5]); set_bg(row.cells[5], 'D9D9D9'); cell_text(row.cells[5], 'اكونت العميل')
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[4]); set_border(cell); cell_text(cell, inv.get('العنوان', ''))
    set_border(row.cells[5]); set_bg(row.cells[5], 'D9D9D9'); cell_text(row.cells[5], 'العنوان')
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[4]); set_border(cell); cell_text(cell, inv.get('الرقم', ''))
    set_border(row.cells[5]); set_bg(row.cells[5], 'D9D9D9'); cell_text(row.cells[5], 'الرقم')
    row = table.add_row()
    set_border(row.cells[0]); cell_text(row.cells[0], inv.get('الألوان', ''))
    cell = row.cells[1].merge(row.cells[2]); set_border(cell); set_bg(cell, 'D9D9D9'); cell_text(cell, 'الألوان')
    set_border(row.cells[3]); cell_text(row.cells[3], inv.get('العدد', ''))
    set_border(row.cells[4]); set_bg(row.cells[4], 'D9D9D9'); cell_text(row.cells[4], 'العدد')
    set_border(row.cells[5]); set_bg(row.cells[5], 'D9D9D9')
    row = table.add_row()
    set_border(row.cells[0]); cell_text(row.cells[0], inv.get('المقاس', ''))
    cell = row.cells[1].merge(row.cells[2]); set_border(cell); set_bg(cell, 'D9D9D9'); cell_text(cell, 'المقاس')
    set_border(row.cells[3]); cell_text(row.cells[3], inv.get('نوع المنتج', ''))
    set_border(row.cells[4]); set_bg(row.cells[4], 'D9D9D9'); cell_text(row.cells[4], 'نوع المنتج')
    set_border(row.cells[5]); set_bg(row.cells[5], 'D9D9D9')
    row = table.add_row()
    set_border(row.cells[0]); cell_text(row.cells[0], inv.get('الاجمالي', ''))
    set_border(row.cells[1]); set_bg(row.cells[1], 'D9D9D9'); cell_text(row.cells[1], 'الاجمالي')
    set_border(row.cells[2]); cell_text(row.cells[2], inv.get('مصاريف الشحن', ''))
    set_border(row.cells[3]); set_bg(row.cells[3], 'D9D9D9'); cell_text(row.cells[3], 'مصاريف الشحن')
    set_border(row.cells[4]); cell_text(row.cells[4], inv.get('سعر المنتج', ''))
    set_border(row.cells[5]); set_bg(row.cells[5], 'D9D9D9'); cell_text(row.cells[5], 'سعر المنتج')
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[5]); set_border(cell); set_bg(cell, 'FFF2CC')
    cell_text(cell, NOTE, bold=False, size=8)

# ====== M&M ======
def make_mm_invoice(doc, inv):
    NOTE = 'يجب دفع خدمة مصاريف الشحن لمندوب شركة الشحن فى حالة الشراء او عدم الشراء نظرا لتكلفة مصاريف الانتقال'
    today = datetime.now().strftime('%d / %m / %Y')
    table = doc.add_table(rows=0, cols=6)
    table.style = 'Table Grid'
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[5])
    set_border(cell); set_bg(cell, '000000')
    cell_text(cell, 'M&M', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color='FFFFFF')
    row = table.add_row()
    set_border(row.cells[0]); cell_text(row.cells[0], today)
    set_border(row.cells[1]); set_bg(row.cells[1], 'D9D9D9'); cell_text(row.cells[1], 'التاريخ')
    set_border(row.cells[2]); cell_text(row.cells[2], inv.get('اسم الاكونت', ''))
    set_border(row.cells[3]); set_bg(row.cells[3], 'D9D9D9'); cell_text(row.cells[3], 'اسم الاكونت')
    set_border(row.cells[4]); cell_text(row.cells[4], inv.get('IT', ''))
    set_border(row.cells[5]); set_bg(row.cells[5], 'D9D9D9'); cell_text(row.cells[5], 'IT')
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[4]); set_border(cell); cell_text(cell, inv.get('الاسم', ''))
    set_border(row.cells[5]); set_bg(row.cells[5], 'D9D9D9'); cell_text(row.cells[5], 'الاسم')
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[4]); set_border(cell); cell_text(cell, inv.get('العنوان', ''))
    set_border(row.cells[5]); set_bg(row.cells[5], 'D9D9D9'); cell_text(row.cells[5], 'العنوان بالتفصيل')
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[4]); set_border(cell); cell_text(cell, inv.get('موبايل 1', ''))
    set_border(row.cells[5]); set_bg(row.cells[5], 'D9D9D9'); cell_text(row.cells[5], 'رقم الموبايل 1')
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[4]); set_border(cell); cell_text(cell, inv.get('موبايل 2', ''))
    set_border(row.cells[5]); set_bg(row.cells[5], 'D9D9D9'); cell_text(row.cells[5], 'رقم الموبايل 2')
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[4]); set_border(cell); cell_text(cell, inv.get('المنتج', ''))
    set_border(row.cells[5]); set_bg(row.cells[5], 'D9D9D9'); cell_text(row.cells[5], 'المنتج 1')
    row = table.add_row()
    set_border(row.cells[0]); cell_text(row.cells[0], inv.get('الالوان', ''))
    cell = row.cells[1].merge(row.cells[2]); set_border(cell); set_bg(cell, 'D9D9D9'); cell_text(cell, 'الالوان')
    set_border(row.cells[3]); cell_text(row.cells[3], inv.get('المقاس', ''))
    set_border(row.cells[4]); set_bg(row.cells[4], 'D9D9D9'); cell_text(row.cells[4], 'المقاس')
    set_border(row.cells[5]); set_bg(row.cells[5], 'D9D9D9')
    row = table.add_row()
    set_border(row.cells[0]); cell_text(row.cells[0], inv.get('الكمية', ''))
    set_border(row.cells[1]); set_bg(row.cells[1], 'D9D9D9'); cell_text(row.cells[1], 'الكمية')
    set_border(row.cells[2]); cell_text(row.cells[2], inv.get('السعر', ''))
    set_border(row.cells[3]); set_bg(row.cells[3], 'D9D9D9'); cell_text(row.cells[3], 'السعر')
    set_border(row.cells[4]); cell_text(row.cells[4], inv.get('الشحن', ''))
    set_border(row.cells[5]); set_bg(row.cells[5], 'D9D9D9'); cell_text(row.cells[5], 'الشحن')
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[4]); set_border(cell); cell_text(cell, inv.get('الاجمالي', ''), bold=True, size=11)
    set_border(row.cells[5]); set_bg(row.cells[5], 'D9D9D9'); cell_text(row.cells[5], 'الاجمالي')
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[5]); set_border(cell); set_bg(cell, 'FFF2CC')
    cell_text(cell, NOTE, bold=False, size=8)

def parse_invoices(text, fields):
    blocks = re.split(r'\n\s*---+\s*\n', text)
    invoices = []
    for block in blocks:
        block = block.strip()
        if not block: continue
        inv = {}
        for line in block.splitlines():
            line = line.strip()
            for field in fields:
                if line.startswith(field + ':'):
                    inv[field] = line[len(field)+1:].strip()
                    break
        if inv: invoices.append(inv)
    return invoices

def create_word(invoices, template_type):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21); section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Cm(1)
    for i, inv in enumerate(invoices):
        if template_type == 'dragon': make_dragons_invoice(doc, inv)
        elif template_type == 'fashion': make_nf_invoice(doc, inv)
        elif template_type == 'mm': make_mm_invoice(doc, inv)
        if i < len(invoices) - 1: doc.add_paragraph()
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

HELP_MESSAGES = {
    'dragon': "✅ اخترت Dragon.s!\n\nابعت /invoice أو ملف .txt بالبيانات:\n\nالاسم: محمد\nالرقم: 01012345678\nالعنوان: القاهرة\naاللون: ازرق\nالعرض: ترنج\nالمقاس: لارج\nالعدد: 1\nالاجمالي: 300",
    'fashion': "✅ اخترت New Fashion!\n\nابعت /invoice أو ملف .txt بالبيانات:\n\naسم العميل: محمد\nاكونت العميل: 01012345678\nرقم الفاتورة: 001\nاسم الصفحة: نيو فاشون\nاسم IT: أحمد\nالعنوان: القاهرة\nالرقم: 01012345678\nالألوان: أحمر\nالعدد: 2\nالمقاس: XL\nنوع المنتج: تيشيرت\nسعر المنتج: 200\nمصاريف الشحن: 30\nالاجمالي: 430",
    'mm': "✅ اخترت M&M!\n\nابعت /invoice أو ملف .txt بالبيانات:\n\naسم الاكونت: محمد\nIT: سارة\nالاسم: محمود\nالعنوان: القاهرة\nموبايل 1: 01012345678\nموبايل 2: 01098765432\nالمنتج: جينز\nالالوان: أزرق\nالمقاس: XL\nالكمية: 2\nالسعر: 350\nالشحن: 40\nالاجمالي: 390",
}

async def process_invoices(update, text, template_type):
    fields = TEMPLATES[template_type]
    invoices = parse_invoices(text, fields)
    if not invoices:
        await update.message.reply_text("❌ مش لاقي فواتير! تأكد من الشكل الصح")
        return
    await update.message.reply_text(f"⏳ بشتغل على {len(invoices)} فاتورة...")
    try:
        buf = create_word(invoices, template_type)
        await update.message.reply_document(
            document=buf,
            filename=f"فواتير_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            caption=f"✅ تم إنشاء *{len(invoices)} فاتورة* بنجاح!",
            parse_mode='Markdown'
        )
    except Exception as e:
        await update.message.reply_text(f"❌ خطأ: {e}")

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    if user_id in authorized_users:
        await update.message.reply_text(
            "أهلاً! 👋\n\nابعتلي اسم شكل الفاتورة:\n• *dragon*\n• *fashion*\n• *mm*",
            parse_mode='Markdown'
        )
        return WAITING_TEMPLATE_CHOICE
    else:
        await update.message.reply_text("أهلاً! 👋\n\n🔐 البوت ده خاص\nابعتلي كود الدخول بتاعك:")
        return WAITING_CODE

async def check_code(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    code = update.message.text.strip()
    if code in valid_codes:
        authorized_users[user_id] = code
        save_data(valid_codes, authorized_users)
        await update.message.reply_text(
            "✅ الكود صح!\n\nابعتلي اسم شكل الفاتورة:\n• *dragon*\n• *fashion*\n• *mm*",
            parse_mode='Markdown'
        )
        return WAITING_TEMPLATE_CHOICE
    else:
        await update.message.reply_text("❌ الكود غلط! جرب تاني")
        return WAITING_CODE

async def choose_template(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    choice = update.message.text.strip().lower()
    if choice in TEMPLATES:
        user_templates[user_id] = choice
        await update.message.reply_text(HELP_MESSAGES[choice])
        return ConversationHandler.END
    else:
        await update.message.reply_text(
            "❌ مش عارف الشكل ده!\n\nاكتب:\n• *dragon*\n• *fashion*\n• *mm*",
            parse_mode='Markdown'
        )
        return WAITING_TEMPLATE_CHOICE

async def invoice_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    text = update.message.text or ""
    if user_id not in authorized_users:
        await update.message.reply_text("❌ محتاج كود دخول! ابعت /start")
        return
    template_type = user_templates.get(user_id)
    if not template_type:
        await update.message.reply_text("❌ ابعت /start الأول عشان تختار شكل الفاتورة")
        return
    text = re.sub(r'^/invoice\s*', '', text.strip(), flags=re.IGNORECASE)
    if len(text.strip()) < 5:
        await update.message.reply_text(HELP_MESSAGES.get(template_type, "ابعت البيانات مع /invoice"))
        return
    await process_invoices(update, text, template_type)

async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    if user_id not in authorized_users:
        await update.message.reply_text("❌ محتاج كود دخول! ابعت /start")
        return
    template_type = user_templates.get(user_id)
    if not template_type:
        await update.message.reply_text("❌ ابعت /start الأول عشان تختار شكل الفاتورة")
        return
    doc = update.message.document
    if not doc.file_name.endswith('.txt'):
        await update.message.reply_text("❌ ابعتلي ملف .txt بس")
        return
    file = await context.bot.get_file(doc.file_id)
    buf = io.BytesIO()
    await file.download_to_memory(buf)
    try:
        text = buf.getvalue().decode('utf-8')
    except:
        text = buf.getvalue().decode('windows-1256', errors='ignore')
    await process_invoices(update, text, template_type)

async def addcode(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID: return
    if not context.args: await update.message.reply_text("استخدام: /addcode الكود"); return
    code = context.args[0]
    valid_codes[code] = datetime.now().strftime('%Y-%m-%d')
    save_data(valid_codes, authorized_users)
    await update.message.reply_text(f"✅ تم إضافة الكود: `{code}`", parse_mode='Markdown')

async def removecode(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID: return
    if not context.args: await update.message.reply_text("استخدام: /removecode الكود"); return
    code = context.args[0]
    if code in valid_codes:
        del valid_codes[code]
        to_remove = [uid for uid, c in authorized_users.items() if c == code]
        for uid in to_remove: del authorized_users[uid]
        save_data(valid_codes, authorized_users)
        await update.message.reply_text(f"✅ تم حذف الكود: `{code}`", parse_mode='Markdown')
    else:
        await update.message.reply_text("❌ الكود مش موجود")

async def listcodes(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID: return
    if not valid_codes: await update.message.reply_text("مفيش أكواد"); return
    text = "📋 *الأكواد:*\n\n"
    for code, date in valid_codes.items():
        users = [uid for uid, c in authorized_users.items() if c == code]
        text += f"• `{code}` - {date} - {len(users)} مستخدم\n"
    await update.message.reply_text(text, parse_mode='Markdown')

async def listusers(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID: return
    if not authorized_users: await update.message.reply_text("مفيش مستخدمين"); return
    text = f"👥 *المستخدمين: {len(authorized_users)}*\n\n"
    for uid, code in authorized_users.items():
        template = user_templates.get(uid, 'مش محدد')
        text += f"• ID: `{uid}` - كود: `{code}` - شكل: {template}\n"
    await update.message.reply_text(text, parse_mode='Markdown')

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    if user_id not in authorized_users:
        await update.message.reply_text("ابعت /start للبدء")
    else:
        await update.message.reply_text("ابعت /invoice لعمل فواتير أو /start لتغيير الشكل 🙂")

def main():
    app = ApplicationBuilder().token(BOT_TOKEN).build()
    conv = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            WAITING_CODE: [MessageHandler(filters.TEXT & ~filters.COMMAND, check_code)],
            WAITING_TEMPLATE_CHOICE: [MessageHandler(filters.TEXT & ~filters.COMMAND, choose_template)],
        },
        fallbacks=[]
    )
    app.add_handler(conv)
    app.add_handler(CommandHandler("invoice", invoice_cmd))
    app.add_handler(CommandHandler("addcode", addcode))
    app.add_handler(CommandHandler("removecode", removecode))
    app.add_handler(CommandHandler("listcodes", listcodes))
    app.add_handler(CommandHandler("listusers", listusers))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_file))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    print("✅ البوت شغال!")
    app.run_polling()

if __name__ == '__main__':
    main()
